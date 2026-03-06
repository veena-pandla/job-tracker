"""
Resume Builder — generates a tailored ATS-friendly Word document for each job.
Uses profile from config.py + AI-tailored sections from ai_engine.py.
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import PROFILE


def _add_heading(doc, text, level=1, color=(0, 0, 0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11 if level == 1 else 10)
    run.font.color.rgb = RGBColor(*color)
    # Add a bottom border line
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_bullet(doc, text, indent=True):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    if indent:
        p.paragraph_format.left_indent = Inches(0.2)
    return p


def _set_margins(doc, margin=Inches(0.75)):
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def build_resume_docx(job: dict, tailored: dict) -> bytes:
    """
    Build a tailored resume Word document.
    Keeps ALL original profile content unchanged.
    Only adds 1-2 new ATS bullets to the most recent job and reorders skills.

    Args:
        job: the job dict (title, company, etc.)
        tailored: {"priority_skills": [...], "extra_bullets": [...]}
    """
    doc = Document()
    _set_margins(doc)

    # Remove default spacing
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── HEADER ────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(PROFILE.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(16)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = []
    if PROFILE.get("email"):
        contact_parts.append(PROFILE["email"])
    if PROFILE.get("phone") and PROFILE["phone"] != "your_phone_number":
        contact_parts.append(PROFILE["phone"])
    if PROFILE.get("linkedin"):
        contact_parts.append(PROFILE["linkedin"])
    if PROFILE.get("location"):
        contact_parts.append(PROFILE["location"])
    contact_run = contact_para.add_run("  |  ".join(contact_parts))
    contact_run.font.size = Pt(9)

    doc.add_paragraph()

    # ── PROFESSIONAL SUMMARY — original, unchanged ─────────────────────────────
    _add_heading(doc, "Professional Summary")
    summary = PROFILE.get("summary", "")
    p = doc.add_paragraph()
    run = p.add_run(summary)
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph()

    # ── SKILLS — priority skills first, then remaining originals ───────────────
    _add_heading(doc, "Technical Skills")
    all_skills = PROFILE.get("skills", [])
    priority = tailored.get("priority_skills", [])
    # Put priority skills first (preserving original casing), then the rest
    priority_set = {s.lower() for s in priority}
    ordered_skills = [s for s in all_skills if s.lower() in priority_set] + \
                     [s for s in all_skills if s.lower() not in priority_set]
    skills_para = doc.add_paragraph()
    skills_run = skills_para.add_run("  •  ".join(ordered_skills))
    skills_run.font.size = Pt(10)

    doc.add_paragraph()

    # ── WORK EXPERIENCE — all original bullets, + 1-2 new ones for first job ──
    _add_heading(doc, "Work Experience")

    extra_bullets = tailored.get("extra_bullets", [])

    for i, exp in enumerate(PROFILE.get("experience", [])):
        # Job title line
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{exp['title']}")
        title_run.bold = True
        title_run.font.size = Pt(10)
        company_run = title_para.add_run(f"  —  {exp['company']}  |  {exp['duration']}")
        company_run.font.size = Pt(10)
        company_run.italic = True

        # All original bullets, unchanged
        for bullet in exp.get("bullets", []):
            _add_bullet(doc, bullet)

        # For most recent job only: append new ATS-targeted bullets at the end
        if i == 0 and extra_bullets:
            for eb in extra_bullets:
                _add_bullet(doc, eb)

        doc.add_paragraph()

    # ── EDUCATION — original, unchanged ───────────────────────────────────────
    _add_heading(doc, "Education")
    for edu in PROFILE.get("education", []):
        edu_para = doc.add_paragraph()
        deg_run = edu_para.add_run(f"{edu['degree']}")
        deg_run.bold = True
        deg_run.font.size = Pt(10)
        school_run = edu_para.add_run(f"  —  {edu['school']}  |  {edu.get('year', '')}")
        school_run.font.size = Pt(10)
        school_run.italic = True

    doc.add_paragraph()

    # ── PROJECTS — all original, unchanged ────────────────────────────────────
    projects = PROFILE.get("projects", [])
    if projects:
        _add_heading(doc, "Key Projects")
        for proj in projects:
            proj_para = doc.add_paragraph()
            proj_run = proj_para.add_run(f"{proj['name']}: ")
            proj_run.bold = True
            proj_run.font.size = Pt(10)
            desc_run = proj_para.add_run(proj.get("description", ""))
            desc_run.font.size = Pt(10)

    # ── Save to bytes ─────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
